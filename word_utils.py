
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.shared import qn
from PIL import Image, UnidentifiedImageError

PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]

def set_font_calibri(run, size=11):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    rFonts = run.font.element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')

def replace_placeholder_in_paragraph(paragraph, placeholder, value):
    if placeholder in paragraph.text:
        # Store original runs and their properties
        original_runs_data = []
        for run in paragraph.runs:
            original_runs_data.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None
            })
        
        # Clear the paragraph and rebuild it
        paragraph.clear()
        for run_data in original_runs_data:
            parts = run_data['text'].split(placeholder)
            for i, part in enumerate(parts):
                if part:
                    new_run = paragraph.add_run(part)
                    new_run.bold = run_data['bold']
                    new_run.italic = run_data['italic']
                    new_run.underline = run_data['underline']
                    set_font_calibri(new_run, 11)
                
                if i < len(parts) - 1:
                    value_run = paragraph.add_run(str(value))
                    value_run.bold = run_data['bold']
                    value_run.italic = run_data['italic']
                    value_run.underline = run_data['underline']
                    set_font_calibri(value_run, 11)

def replace_all_placeholders(doc, campos):
    """Substitui todos os placeholders no documento"""
    for key, value in campos.items():
        placeholder = f'{{{{{key.lower()}}}}}'
        
        # Substituir em parágrafos
        for paragraph in doc.paragraphs:
            replace_placeholder_in_paragraph(paragraph, placeholder, value)
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, placeholder, value)
        
        # Substituir em cabeçalhos e rodapés
        for section in doc.sections:
            # Cabeçalho
            for paragraph in section.header.paragraphs:
                replace_placeholder_in_paragraph(paragraph, placeholder, value)
            
            # Rodapé
            for paragraph in section.footer.paragraphs:
                replace_placeholder_in_paragraph(paragraph, placeholder, value)

def aplicar_estilo(run, tamanho, negrito=False):
    run.font.name = "Arial"
    run.font.size = Pt(tamanho)
    run.bold = negrito

def inserir_conteudo(modelo_path, conteudo, output_path, campos=None):
    doc = Document(modelo_path)
    contador_imagens = 0
    conteudo_processado = False
    paragrafo_insercao_index = None

    # --- Substituir placeholders com dados do formulário ---
    if campos:
        replace_all_placeholders(doc, campos)
    # --- Fim da substituição de placeholders ---

    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("Marca '{{start_here}}' não encontrada no modelo.")
        return contador_imagens

    conteudo_invertido = list(reversed(conteudo))

    for item in conteudo_invertido:
        if isinstance(item, str):
            titulo = item.replace("»", "").strip() + ":"
            nivel = item.count("»")

            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(titulo)

            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, negrito=True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = 'Heading 1'
            elif nivel == 1:
                p.style = 'Heading 2'
            elif nivel == 2:
                p.style = 'Heading 3'
            else:
                aplicar_estilo(run, 12, negrito=True)

            conteudo_processado = True

        elif isinstance(item, dict):
            if 'image_path' in item:
                image_path = item['image_path']
                if os.path.exists(image_path) and os.path.getsize(image_path) > 0:
                    try:
                        # Inserção da imagem (sem legenda)
                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        with Image.open(image_path) as img:
                            largura_original, altura_original = img.size
                            altura_desejada_cm = 10
                            proporcao = altura_desejada_cm / altura_original * 2.54
                            largura_proporcional_cm = largura_original * proporcao / 2.54

                            run = p.add_run()
                            run.add_picture(
                                image_path,
                                width=Cm(largura_proporcional_cm),
                                height=Cm(altura_desejada_cm)
                            )
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            contador_imagens += 1
                            conteudo_processado = True

                    except UnidentifiedImageError:
                        print(f"Erro: Formato de imagem não reconhecido: {image_path}")
                    except Exception as e:
                        print(f"Erro ao inserir imagem '{image_path}': {e}")
                else:
                    print(f"Erro: Arquivo de imagem inválido: {image_path}")

            elif 'quebra_pagina' in item:
                doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)
                conteudo_processado = False

    doc.save(output_path)
    return contador_imagens
